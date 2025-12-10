"""
Document generation service for creating Word documents.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict
import io
import zipfile
import os
from utils.logger import log_step, log_success, log_debug


class DocumentService:
    """Generate Word documents for quizzes and assignments."""
    
    @staticmethod
    def create_quiz_document(quizzes: List[Dict]) -> bytes:
        """
        Create a Word document containing all quizzes.
        
        Args:
            quizzes: List of quiz dictionaries
            
        Returns:
            Bytes of the Word document
        """
        log_step("Creating Quiz Document", f"Generating document for {len(quizzes)} quiz(es)")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Generated Quizzes', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for quiz in quizzes:
            # Quiz header
            doc.add_heading(f"Quiz {quiz['quiz_number']}", level=1)
            doc.add_paragraph(f"Total Marks: {quiz['total_marks']}")
            doc.add_paragraph("")
            
            for q in quiz['questions']:
                # Question
                q_para = doc.add_paragraph()
                q_para.add_run(f"Q{q['question_number']}. ").bold = True
                q_para.add_run(q['question_text'])
                
                # Question metadata
                meta = doc.add_paragraph()
                meta.add_run(f"[{q['question_type'].upper()} | {q['marks']} marks | {q['difficulty_level']}]")
                meta.runs[0].font.size = Pt(9)
                meta.runs[0].font.italic = True
                
                # Options for MCQ
                if q.get('options'):
                    for i, opt in enumerate(q['options']):
                        option_letter = chr(65 + i)  # A, B, C, D
                        doc.add_paragraph(f"    {option_letter}. {opt}")
                
                doc.add_paragraph("")  # Spacing
            
            # Answer Key Section
            doc.add_heading("Answer Key", level=2)
            for q in quiz['questions']:
                ans_para = doc.add_paragraph()
                ans_para.add_run(f"Q{q['question_number']}: ").bold = True
                ans_para.add_run(q['correct_answer'])
                if q.get('explanation'):
                    exp_para = doc.add_paragraph(f"   Explanation: {q['explanation']}")
                    exp_para.runs[0].font.size = Pt(9)
            
            doc.add_page_break()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        log_success(f"Quiz document created successfully")
        return buffer.getvalue()
    
    @staticmethod
    def create_assignment_document(assignments: List[Dict]) -> bytes:
        """
        Create a Word document containing all assignments.
        
        Args:
            assignments: List of assignment dictionaries
            
        Returns:
            Bytes of the Word document
        """
        log_step("Creating Assignment Document", f"Generating document for {len(assignments)} assignment(s)")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Generated Assignments', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for assignment in assignments:
            # Assignment header
            doc.add_heading(f"Assignment {assignment['assignment_number']}", level=1)
            doc.add_paragraph(f"Total Marks: {assignment['total_marks']}")
            doc.add_paragraph("")
            
            for q in assignment['questions']:
                # Question
                q_para = doc.add_paragraph()
                q_para.add_run(f"Q{q['question_number']}. ").bold = True
                q_para.add_run(q['question_text'])
                
                # Question metadata
                meta = doc.add_paragraph()
                meta_text = f"[{q['marks']} marks | {q['difficulty_level']}"
                if q.get('expected_length'):
                    meta_text += f" | Expected: {q['expected_length']}"
                meta_text += "]"
                meta.add_run(meta_text)
                meta.runs[0].font.size = Pt(9)
                meta.runs[0].font.italic = True
                
                # Key points if available
                if q.get('key_points'):
                    doc.add_paragraph("Key points to cover:")
                    for point in q['key_points']:
                        doc.add_paragraph(f"    • {point}")
                
                doc.add_paragraph("")  # Spacing
            
            doc.add_page_break()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        log_success(f"Assignment document created successfully")
        return buffer.getvalue()
    
    @staticmethod
    def create_combined_document(quizzes: List[Dict], assignments: List[Dict]) -> bytes:
        """
        Create a combined Word document with both quizzes and assignments.
        
        Args:
            quizzes: List of quiz dictionaries
            assignments: List of assignment dictionaries
            
        Returns:
            Bytes of the Word document
        """
        log_step("Creating Combined Document", f"{len(quizzes)} quiz(es) + {len(assignments)} assignment(s)")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Quiz & Assignment Package', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        
        # Table of Contents
        doc.add_heading("Contents", level=1)
        if quizzes:
            for quiz in quizzes:
                doc.add_paragraph(f"• Quiz {quiz['quiz_number']} ({quiz['total_marks']} marks)")
        if assignments:
            for assignment in assignments:
                doc.add_paragraph(f"• Assignment {assignment['assignment_number']} ({assignment['total_marks']} marks)")
        
        doc.add_page_break()
        
        # Quizzes Section
        if quizzes:
            doc.add_heading("SECTION A: QUIZZES", level=1)
            
            for quiz in quizzes:
                doc.add_heading(f"Quiz {quiz['quiz_number']}", level=2)
                doc.add_paragraph(f"Total Marks: {quiz['total_marks']}")
                doc.add_paragraph("")
                
                for q in quiz['questions']:
                    q_para = doc.add_paragraph()
                    q_para.add_run(f"Q{q['question_number']}. ").bold = True
                    q_para.add_run(q['question_text'])
                    
                    meta = doc.add_paragraph()
                    meta.add_run(f"[{q['question_type'].upper()} | {q['marks']} marks | {q['difficulty_level']}]")
                    meta.runs[0].font.size = Pt(9)
                    meta.runs[0].font.italic = True
                    
                    if q.get('options'):
                        for i, opt in enumerate(q['options']):
                            option_letter = chr(65 + i)
                            doc.add_paragraph(f"    {option_letter}. {opt}")
                    
                    doc.add_paragraph("")
                
                # Answer key
                doc.add_heading("Answer Key", level=3)
                for q in quiz['questions']:
                    ans_para = doc.add_paragraph()
                    ans_para.add_run(f"Q{q['question_number']}: ").bold = True
                    ans_para.add_run(q['correct_answer'])
                
                doc.add_page_break()
        
        # Assignments Section
        if assignments:
            doc.add_heading("SECTION B: ASSIGNMENTS", level=1)
            
            for assignment in assignments:
                doc.add_heading(f"Assignment {assignment['assignment_number']}", level=2)
                doc.add_paragraph(f"Total Marks: {assignment['total_marks']}")
                doc.add_paragraph("")
                
                for q in assignment['questions']:
                    q_para = doc.add_paragraph()
                    q_para.add_run(f"Q{q['question_number']}. ").bold = True
                    q_para.add_run(q['question_text'])
                    
                    meta = doc.add_paragraph()
                    meta_text = f"[{q['marks']} marks | {q['difficulty_level']}"
                    if q.get('expected_length'):
                        meta_text += f" | Expected: {q['expected_length']}"
                    meta_text += "]"
                    meta.add_run(meta_text)
                    meta.runs[0].font.size = Pt(9)
                    meta.runs[0].font.italic = True
                    
                    if q.get('key_points'):
                        doc.add_paragraph("Key points to cover:")
                        for point in q['key_points']:
                            doc.add_paragraph(f"    • {point}")
                    
                    doc.add_paragraph("")
                
                doc.add_page_break()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        log_success("Combined document created successfully")
        return buffer.getvalue()
    
    @staticmethod
    def create_student_answer_document(filename: str, answers: List[Dict], quiz_answers: List[Dict] = None, raw_text: str = None, extraction_stats: List[Dict] = None) -> bytes:
        """
        Create a Word document for a single student's extracted answers.
        
        Args:
            filename: Original filename (student identifier)
            answers: List of extracted answers
            quiz_answers: List of quiz/MCQ answers (optional)
            raw_text: Complete raw extracted text (optional)
            extraction_stats: Extraction statistics (optional)
            
        Returns:
            Bytes of the Word document
        """
        log_step("Creating Student Answer Document", f"File: {filename}")
        
        doc = Document()
        
        # Title with student/file identifier
        student_name = os.path.splitext(filename)[0]  # Remove extension
        title = doc.add_heading(f'Extracted Answers', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # File info
        doc.add_paragraph(f"Source File: {filename}")
        doc.add_paragraph(f"Total Answers Extracted: {len(answers)}")
        if quiz_answers:
            doc.add_paragraph(f"Quiz/MCQ Answers: {len(quiz_answers)}")
        
        # Extraction stats
        if extraction_stats:
            total_words = sum(s.get('words_extracted', 0) for s in extraction_stats if isinstance(s.get('words_extracted'), int))
            if total_words:
                doc.add_paragraph(f"Approximate Words Extracted: {total_words}")
        
        doc.add_paragraph("")
        
        # RAW TEXT SECTION FIRST
        if raw_text:
            doc.add_heading("Complete Extracted Text (Raw)", level=1)
            doc.add_paragraph("This is all the text extracted from the document before structuring:")
            doc.add_paragraph("")
            
            # Add raw text with preserved formatting
            raw_para = doc.add_paragraph()
            raw_para.add_run(raw_text)
            raw_para.runs[0].font.size = Pt(10)
            
            doc.add_page_break()
        
        doc.add_heading("─" * 40, level=2)
        
        # Long/Short Answers Section
        if answers:
            doc.add_heading("Answers", level=1)
            
            # Sort answers by answer number if possible
            sorted_answers = sorted(answers, key=lambda x: str(x.get('answer_number', '0')))
            
            for answer in sorted_answers:
                ans_num = answer.get('answer_number', '?')
                ans_type = answer.get('answer_type', 'unknown')
                content = answer.get('content', 'No content extracted')
                confidence = answer.get('confidence', 'N/A')
                pages = answer.get('pages', [])
                
                # Answer header
                header = doc.add_heading(f"Answer {ans_num}", level=2)
                
                # Metadata
                meta = doc.add_paragraph()
                meta.add_run(f"Type: {ans_type} | Confidence: {confidence}")
                if pages:
                    meta.add_run(f" | Pages: {', '.join(map(str, pages))}")
                meta.runs[0].font.size = Pt(9)
                meta.runs[0].font.italic = True
                
                # Answer content
                doc.add_paragraph("")
                content_para = doc.add_paragraph(content)
                
                # Separator
                doc.add_paragraph("")
                doc.add_paragraph("─" * 50)
                doc.add_paragraph("")
        
        # Quiz/MCQ Answers Section
        if quiz_answers:
            doc.add_page_break()
            doc.add_heading("Quiz / MCQ Answers", level=1)
            
            for qa in quiz_answers:
                q_num = qa.get('question_number', '?')
                answer = qa.get('answer', 'N/A')
                confidence = qa.get('confidence', 'N/A')
                
                qa_para = doc.add_paragraph()
                qa_para.add_run(f"Question {q_num}: ").bold = True
                qa_para.add_run(answer)
                qa_para.add_run(f"  (Confidence: {confidence})")
                qa_para.runs[-1].font.size = Pt(9)
                qa_para.runs[-1].font.italic = True
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        log_success(f"Student answer document created: {filename}")
        return buffer.getvalue()
    
    @staticmethod
    def create_all_student_documents(files_data: List[Dict]) -> bytes:
        """
        Create a ZIP file containing Word documents for each student/file.
        
        Args:
            files_data: List of dicts with 'filename', 'answers', 'quiz_answers', 'raw_text', 'extraction_stats'
            
        Returns:
            Bytes of the ZIP file
        """
        log_step("Creating ZIP of Student Documents", f"{len(files_data)} files")
        
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for file_data in files_data:
                filename = file_data.get('filename', 'unknown.pdf')
                answers = file_data.get('answers', [])
                quiz_answers = file_data.get('quiz_answers', [])
                raw_text = file_data.get('raw_text', '')
                extraction_stats = file_data.get('extraction_stats', [])
                
                # Create Word doc for this student
                doc_bytes = DocumentService.create_student_answer_document(
                    filename, answers, quiz_answers, raw_text, extraction_stats
                )
                
                # Create output filename
                base_name = os.path.splitext(filename)[0]
                output_name = f"{base_name}_answers.docx"
                
                # Add to ZIP
                zip_file.writestr(output_name, doc_bytes)
                log_debug(f"Added to ZIP: {output_name}")
        
        zip_buffer.seek(0)
        log_success(f"Created ZIP with {len(files_data)} documents")
        return zip_buffer.getvalue()

